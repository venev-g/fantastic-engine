from typing import Dict, List
from langchain.document_loaders import UnstructuredFileLoader
from langchain_core.documents import Document
from langchain_text_splitters import (
    RecursiveCharacterTextSplitter,
    MarkdownHeaderTextSplitter,
    HTMLHeaderTextSplitter
)
import os
from bs4 import BeautifulSoup
from unstructured.partition.html import partition_html
from unstructured.partition.pdf import partition_pdf
import PyPDF2
import tabula
import pandas as pd
import re

class DocumentAnalyzer:
    def __init__(self):
        pass

    def load_document(self, file_path: str) -> Dict:
        loader = UnstructuredFileLoader(file_path)
        raw_documents = loader.load()
        return {"elements": raw_documents}

    def _process_pdf(self, file_path: str) -> Dict:
        """Process PDF documents."""
        document_structure = {
            'metadata': {'source': file_path, 'type': 'pdf'},
            'elements': []
        }

        try:
            # Check if poppler is available
            import shutil
            if not shutil.which('pdftotext'):
                print("Warning: poppler-utils not found in PATH. Installing...")
                import subprocess
                subprocess.run(["apt-get", "update"], check=True)
                subprocess.run(["apt-get", "install", "-y", "poppler-utils"], check=True)
                print("poppler-utils installed successfully")

            # Read PDF and extract text
            elements = partition_pdf(
                filename=file_path,
                extract_images=True,
                infer_table_structure=True
            )

            # Extract tables separately using tabula
            tables = tabula.read_pdf(file_path, pages='all')

            # Process elements and categorize them
            for element in elements:
                elem_type = str(type(element)).lower()
                element_data = {
                    'content': str(element),
                    'type': None
                }

                if 'title' in elem_type or 'heading' in elem_type:
                    element_data['type'] = 'heading'
                elif 'table' in elem_type:
                    element_data['type'] = 'table'
                elif 'image' in elem_type:
                    element_data['type'] = 'image' 
                elif 'text' in elem_type:
                    # Check if it's strikeout or highlighted (would need more PDF-specific analysis)
                    if '~~' in str(element) or '--' in str(element):
                        element_data['type'] = 'strikeout'
                    elif any(marker in str(element) for marker in ['**', '__', '>>']):
                        element_data['type'] = 'highlight'
                    else:
                        element_data['type'] = 'paragraph'

                document_structure['elements'].append(element_data)

            # Add tables from tabula to our elements list
            for i, table in enumerate(tables):
                document_structure['elements'].append({
                    'content': table,
                    'type': 'table',
                    'pandas_table': True,
                    'table_id': i
                })
        except Exception as e:
            print(f"Error processing PDF: {e}")
            document_structure['elements'].append({
                'content': f"Error processing PDF file: {str(e)}",
                'type': 'paragraph',
        })

        return document_structure

    def _process_html(self, file_path: str) -> Dict:
        """Process HTML documents."""
        with open(file_path, 'r', encoding='utf-8') as file:
            html_content = file.read()

        soup = BeautifulSoup(html_content, 'lxml')
        document_structure = {
            'metadata': {'source': file_path, 'type': 'html'},
            'elements': []
        }

        # Extract headings
        for heading_level in range(1, 7):
            for heading in soup.find_all(f'h{heading_level}'):
                document_structure['elements'].append({
                    'content': heading.get_text(),
                    'type': 'heading',
                    'level': heading_level
                })

        # Extract paragraphs
        for para in soup.find_all('p'):
            # Check for highlighted text
            highlighted = para.find_all(['strong', 'b', 'mark', 'em'])
            strikeout = para.find_all('s')

            if highlighted:
                for h in highlighted:
                    document_structure['elements'].append({
                        'content': h.get_text(),
                        'type': 'highlight',
                    })

            if strikeout:
                for s in strikeout:
                    document_structure['elements'].append({
                        'content': s.get_text(),
                        'type': 'strikeout',
                    })

            # Add the full paragraph too
            document_structure['elements'].append({
                'content': para.get_text(),
                'type': 'paragraph',
            })

        # Extract tables
        for table in soup.find_all('table'):
            # Convert HTML table to pandas DataFrame
            table_data = []
            rows = table.find_all('tr')
            for row in rows:
                cols = row.find_all(['td', 'th'])
                cols = [ele.get_text().strip() for ele in cols]
                table_data.append(cols)

            if table_data:
                # Try to create a pandas DataFrame
                try:
                    df = pd.DataFrame(table_data[1:], columns=table_data[0])
                    document_structure['elements'].append({
                        'content': df,
                        'type': 'table',
                        'pandas_table': True
                    })
                except:
                    # Fallback to string representation
                    document_structure['elements'].append({
                        'content': str(table_data),
                        'type': 'table',
                        'pandas_table': False
                    })

        # Extract images
        for img in soup.find_all('img'):
            document_structure['elements'].append({
                'content': img.get('alt', 'Image') + f" (src: {img.get('src', '')})",
                'type': 'image',
            })

        return document_structure

    def _process_text(self, file_path: str) -> Dict:
        """Process plain text or markdown documents."""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()

        document_structure = {
            'metadata': {'source': file_path, 'type': 'text'},
            'elements': []
        }

        # Split by double newlines to separate paragraphs
        paragraphs = content.split('\n\n')

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # Check if it's a heading (starts with # in markdown)
            if para.startswith('#'):
                level = len(re.match(r'^#+', para).group())
                heading_text = para.lstrip('#').strip()
                document_structure['elements'].append({
                    'content': heading_text,
                    'type': 'heading',
                    'level': level
                })
            # Check if it's a table (simple detection for markdown tables)
            elif '|' in para and '-+-' in para.replace('|', '+'):
                document_structure['elements'].append({
                    'content': para,
                    'type': 'table',
                    'pandas_table': False
                })
            # Check for strikeout text (~~text~~ in markdown)
            elif '~~' in para:
                document_structure['elements'].append({
                    'content': para,
                    'type': 'strikeout',
                })
            # Check for highlighted text (** or __ in markdown)
            elif '**' in para or '__' in para:
                document_structure['elements'].append({
                    'content': para,
                    'type': 'highlight',
                })
            # Regular paragraph
            else:
                document_structure['elements'].append({
                    'content': para,
                    'type': 'paragraph',
                })

        return document_structure

    def _process_word(self, file_path: str) -> Dict:
        """
        Process Word documents using python-docx library.
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            Dictionary with structured document content
        """
        try:
            import docx
            from docx.oxml.table import CT_Tbl
            from docx.oxml.text.paragraph import CT_P
            
            document = docx.Document(file_path)
            document_structure = {
                'metadata': {'source': file_path, 'type': 'docx'},
                'elements': []
            }
            
            # Process document body
            for element in document.element.body:
                # Process paragraphs
                if isinstance(element, CT_P):
                    paragraph = docx.Document().add_paragraph()
                    paragraph._p = element
                    text = paragraph.text.strip()
                    
                    if not text:
                        continue
                    
                    # Check if it's likely a heading (based on style)
                    p = document.add_paragraph()
                    p._p = element
                    if hasattr(p, 'style') and p.style and 'heading' in p.style.name.lower():
                        document_structure['elements'].append({
                            'content': text,
                            'type': 'heading',
                            'level': int(p.style.name[-1]) if p.style.name[-1].isdigit() else 1
                        })
                    # Check for highlighted or strikeout text
                    elif '**' in text or '__' in text:
                        document_structure['elements'].append({
                            'content': text,
                            'type': 'highlight'
                        })
                    elif '~~' in text:
                        document_structure['elements'].append({
                            'content': text,
                            'type': 'strikeout'
                        })
                    else:
                        document_structure['elements'].append({
                            'content': text,
                            'type': 'paragraph'
                        })
                
                # Process tables
                elif isinstance(element, CT_Tbl):
                    table = docx.Document().add_table(rows=1, cols=1)
                    table._tbl = element
                    
                    # Convert table to pandas DataFrame
                    data = []
                    headers = []
                    
                    # Get headers from first row
                    if table.rows:
                        for cell in table.rows[0].cells:
                            headers.append(cell.text.strip())
                    
                    # Get data from remaining rows
                    for row in table.rows[1:]:
                        row_data = []
                        for cell in row.cells:
                            row_data.append(cell.text.strip())
                        data.append(row_data)
                    
                    # Create pandas DataFrame if possible
                    try:
                        if headers and data:
                            df = pd.DataFrame(data, columns=headers)
                            document_structure['elements'].append({
                                'content': df,
                                'type': 'table',
                                'pandas_table': True
                            })
                        else:
                            # Create simple text representation for table
                            table_text = "Table content:\n"
                            for row in table.rows:
                                row_text = [cell.text.strip() for cell in row.cells]
                                table_text += " | ".join(row_text) + "\n"
                            document_structure['elements'].append({
                                'content': table_text,
                                'type': 'table',
                                'pandas_table': False
                            })
                    except Exception as e:
                        print(f"Error converting table to DataFrame: {e}")
                        table_text = "Table content (error converting):\n"
                        for row in table.rows:
                            row_text = [cell.text.strip() for cell in row.cells]
                            table_text += " | ".join(row_text) + "\n"
                        document_structure['elements'].append({
                            'content': table_text,
                            'type': 'table',
                            'pandas_table': False
                        })
            
            # Process document properties (metadata)
            try:
                core_properties = document.core_properties
                document_structure['metadata']['title'] = core_properties.title
                document_structure['metadata']['author'] = core_properties.author
                document_structure['metadata']['created'] = str(core_properties.created)
                document_structure['metadata']['modified'] = str(core_properties.modified)
            except Exception as e:
                print(f"Error extracting document properties: {e}")
            
            return document_structure
        
        except ImportError:
            print("python-docx package not found. Please install it with 'pip install python-docx'")
            return {
                'metadata': {'source': file_path, 'type': 'docx'},
                'elements': [{
                    'content': f"python-docx package required for Word processing: {file_path}",
                    'type': 'paragraph',
                }]
            }
        except Exception as e:
            print(f"Error processing Word document: {e}")
            return {
                'metadata': {'source': file_path, 'type': 'docx'},
                'elements': [{
                    'content': f"Error processing Word document: {file_path}. Error: {str(e)}",
                    'type': 'paragraph',
                }]
            }

    def extract_tables_as_triplets(self, document_structure: Dict) -> List[Dict]:
        """
        Extract tables from document and convert to triplets for Neo4j.
        """
        triplets = []
        table_count = 0
        
        try:
            for idx, element in enumerate(document_structure['elements']):
                if element['type'] == 'table':
                    if element.get('pandas_table', False):
                        try:
                            df = element['content']
                            if isinstance(df, pd.DataFrame):
                                # Get table name
                                table_name = f"Table_{table_count}"
                                table_count += 1
                                
                                # Look for the closest heading before this table
                                for j in range(idx-1, -1, -1):
                                    if j < len(document_structure['elements']) and document_structure['elements'][j]['type'] == 'heading':
                                        table_name = document_structure['elements'][j]['content']
                                        break
                                        
                                # Create triplets from table
                                row_count = 0
                                for _, row in df.iterrows():
                                    for col in df.columns:
                                        try:
                                            triplets.append({
                                                'subject': f"{table_name}",
                                                'predicate': str(col),
                                                'object': str(row[col]),
                                                'row_id': row_count
                                            })
                                        except Exception as e:
                                            print(f"Error creating triplet for cell: {e}")
                                    row_count += 1
                        except Exception as e:
                            print(f"Error processing table at index {idx}: {e}")
        except Exception as e:
            print(f"Error in extract_tables_as_triplets: {e}")
            
        return triplets
