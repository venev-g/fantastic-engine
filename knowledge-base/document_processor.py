import os
import re
import PyPDF2
import pandas as pd
import numpy as np
from docx import Document
from bs4 import BeautifulSoup
from PIL import Image
import pytesseract
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Milvus
from langchain.docstore.document import Document as LangchainDocument
from langchain.schema import Document as LCDocument
from langchain_community.retrievers import MilvusRetriever
from langchain_community.graphs import Neo4jGraph
from langchain_core.vectorstores import VectorStoreRetriever
from langchain.chains import RetrievalQA
from langchain_community.llms import LlamaCpp

class DocumentProcessor:
    def __init__(self, model_name="sentence-transformers/all-MiniLM-L6-v2"):
        """Initialize the document processor with a local embedding model."""
        self.embeddings = HuggingFaceEmbeddings(model_name=model_name)
        self.milvus_client = None
        self.neo4j_graph = None
        
    def connect_to_databases(self, milvus_host="localhost", milvus_port=19530, 
                             neo4j_uri="bolt://localhost:7687", neo4j_user="neo4j", neo4j_password="password"):
        """Connect to Milvus and Neo4j databases."""
        # Initialize Neo4j connection
        self.neo4j_graph = Neo4jGraph(
            url=neo4j_uri,
            username=neo4j_user,
            password=neo4j_password
        )
        
        # Milvus connection is handled by the Milvus class when creating collections
        
    def extract_text_from_pdf(self, pdf_path):
        """Extract text content from a PDF file."""
        text = ""
        with open(pdf_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    def extract_text_from_docx(self, docx_path):
        """Extract text, including styles and structure from a DOCX file."""
        doc = Document(docx_path)
        document_data = {
            "headings": [],
            "paragraphs": [],
            "tables": []
        }
        
        # Process headings and paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                document_data["headings"].append({
                    "level": int(paragraph.style.name.replace('Heading', '')),
                    "text": paragraph.text,
                    "position": len(document_data["paragraphs"])
                })
            if paragraph.text.strip():
                # Check for highlighted or strikethrough text
                runs_data = []
                for run in paragraph.runs:
                    runs_data.append({
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": run.underline,
                        "strikethrough": run.strike,
                        "highlight": run.highlight_color is not None
                    })
                document_data["paragraphs"].append({
                    "text": paragraph.text,
                    "runs": runs_data
                })
        
        # Process tables
        for table in doc.tables:
            table_data = []
            header_row = []
            for i, row in enumerate(table.rows):
                row_data = [cell.text for cell in row.cells]
                if i == 0:
                    header_row = row_data
                else:
                    table_data.append(row_data)
            document_data["tables"].append({
                "header": header_row,
                "rows": table_data
            })
            
        return document_data
    
    def extract_text_from_html(self, html_path):
        """Extract text and structure from HTML content."""
        with open(html_path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f, 'html.parser')
        
        # Extract headings
        headings = []
        for i in range(1, 7):
            for heading in soup.find_all(f'h{i}'):
                headings.append({
                    "level": i,
                    "text": heading.get_text().strip()
                })
        
        # Extract paragraphs
        paragraphs = []
        for p in soup.find_all('p'):
            paragraphs.append(p.get_text().strip())
        
        # Extract tables
        tables = []
        for table in soup.find_all('table'):
            header_row = []
            header = table.find('thead')
            if header:
                header_cells = header.find_all(['th', 'td'])
                header_row = [cell.get_text().strip() for cell in header_cells]
            
            rows = []
            tbody = table.find('tbody') or table
            for tr in tbody.find_all('tr'):
                row = [td.get_text().strip() for td in tr.find_all(['td', 'th'])]
                if row and not (len(row) == len(header_row) and row == header_row):
                    rows.append(row)
            
            tables.append({
                "header": header_row,
                "rows": rows
            })
        
        # Handle strikethrough and highlighted text
        strikethrough_text = [s.get_text().strip() for s in soup.find_all(['s', 'strike', 'del'])]
        highlighted_text = [mark.get_text().strip() for mark in soup.find_all('mark')]
        
        return {
            "headings": headings,
            "paragraphs": paragraphs,
            "tables": tables,
            "strikethrough": strikethrough_text,
            "highlighted": highlighted_text
        }
    
    def extract_text_from_image(self, image_path):
        """Extract text from images using OCR."""
        image = Image.open(image_path)
        text = pytesseract.image_to_string(image)
        return text
    
    def analyze_document(self, file_path):
        """
        Analyze a document and extract its content based on file type.
        Returns a structured representation of the document.
        """
        # Determine file type based on extension
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_ext in ['.html', '.htm']:
            return self.extract_text_from_html(file_path)
        elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff']:
            return self.extract_text_from_image(file_path)
        elif file_ext in ['.txt', '.md']:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

    def hierarchical_chunking(self, document_data):
        """
        Implement Title + Content chunking (Hierarchical Chunking).
        Pairs headings with their associated content.
        """
        chunks = []
        
        # Check if document_data is a dictionary (structured) or a string (plain text)
        if isinstance(document_data, dict):
            # For structured documents like DOCX or HTML
            headings = document_data.get("headings", [])
            paragraphs = document_data.get("paragraphs", [])
            
            # Sort headings by position
            headings_sorted = sorted(headings, key=lambda h: h.get("position", 0))
            
            # Group content under headings
            current_heading = {"level": 0, "text": "Document Root", "position": -1}
            current_content = []
            
            for i, paragraph in enumerate(paragraphs):
                # Check if we've reached a new heading
                next_heading_pos = next((h["position"] for h in headings_sorted 
                                         if h["position"] > current_heading["position"]), len(paragraphs))
                
                if i < next_heading_pos:
                    # Still under current heading
                    if isinstance(paragraph, dict):
                        current_content.append(paragraph["text"])
                    else:
                        current_content.append(paragraph)
                else:
                    # Save current chunk and start a new one
                    if current_content:
                        chunks.append({
                            "title": current_heading["text"],
                            "content": "\n".join(current_content),
                            "metadata": {
                                "heading_level": current_heading["level"],
                                "chunk_type": "title_content"
                            }
                        })
                    
                    # Find the new heading
                    current_heading = next((h for h in headings_sorted if h["position"] == i), current_heading)
                    current_content = []
                    
                    # Add the current paragraph
                    if isinstance(paragraph, dict):
                        current_content.append(paragraph["text"])
                    else:
                        current_content.append(paragraph)
            
            # Don't forget the last chunk
            if current_content:
                chunks.append({
                    "title": current_heading["text"],
                    "content": "\n".join(current_content),
                    "metadata": {
                        "heading_level": current_heading["level"],
                        "chunk_type": "title_content"
                    }
                })
                
        else:
            # For plain text documents, attempt to identify headers and content
            text = document_data
            # Simple heuristic: Look for lines that might be headers
            # (short lines that don't end with a period)
            lines = text.split('\n')
            i = 0
            current_title = "Document"
            current_content = []
            
            while i < len(lines):
                line = lines[i].strip()
                # Potential heading detection
                if line and len(line) < 100 and not line.endswith('.') and not line.endswith(','):
                    # Save previous chunk if it exists
                    if current_content:
                        chunks.append({
                            "title": current_title,
                            "content": "\n".join(current_content),
                            "metadata": {
                                "chunk_type": "title_content"
                            }
                        })
                    
                    # Start new chunk
                    current_title = line
                    current_content = []
                else:
                    if line:  # Only add non-empty lines
                        current_content.append(line)
                i += 1
            
            # Don't forget the last chunk
            if current_content:
                chunks.append({
                    "title": current_title,
                    "content": "\n".join(current_content),
                    "metadata": {
                        "chunk_type": "title_content"
                    }
                })
                
        return chunks

    def paragraph_chunking(self, document_data):
        """
        Implement Paragraph-Based Chunking with line breaks preserved.
        """
        chunks = []
        
        if isinstance(document_data, dict):
            paragraphs = document_data.get("paragraphs", [])
            
            for para in paragraphs:
                if isinstance(para, dict):
                    chunks.append({
                        "content": para["text"],
                        "metadata": {
                            "chunk_type": "paragraph",
                            "has_formatting": len(para.get("runs", [])) > 1
                        }
                    })
                else:
                    chunks.append({
                        "content": para,
                        "metadata": {
                            "chunk_type": "paragraph"
                        }
                    })
        else:
            # For plain text, split on double new lines (typical paragraph breaks)
            text = document_data
            paragraphs = re.split(r'\n\s*\n', text)
            
            for para in paragraphs:
                if para.strip():  # Only include non-empty paragraphs
                    chunks.append({
                        "content": para.strip(),
                        "metadata": {
                            "chunk_type": "paragraph"
                        }
                    })
                    
        return chunks

    def table_chunking(self, document_data):
        """
        Process tables and convert them to graph-friendly triplets.
        """
        triplets = []
        
        if isinstance(document_data, dict):
            tables = document_data.get("tables", [])
            
            for table_idx, table in enumerate(tables):
                headers = table.get("header", [])
                rows = table.get("rows", [])
                
                # Generate a table ID
                table_id = f"table_{table_idx}"
                
                # Create triplets representing table structure
                triplets.append({
                    "source": "document",
                    "relation": "has_table",
                    "target": table_id
                })
                
                # Process each row
                for row_idx, row in enumerate(rows):
                    # Create a row entity
                    row_id = f"{table_id}_row_{row_idx}"
                    triplets.append({
                        "source": table_id,
                        "relation": "has_row",
                        "target": row_id
                    })
                    
                    # Create cell entities with their values
                    for col_idx, cell_value in enumerate(row):
                        # Get column name (header) if available, otherwise use index
                        col_name = headers[col_idx] if col_idx < len(headers) else f"column_{col_idx}"
                        
                        triplets.append({
                            "source": row_id,
                            "relation": col_name,
                            "target": cell_value
                        })
        
        return triplets

    def metadata_aware_chunking(self, document_data):
        """
        Create chunks that preserve metadata like highlighting, strikethrough, etc.
        """
        chunks = []
        
        if isinstance(document_data, dict):
            # Process highlighted text
            highlighted = document_data.get("highlighted", [])
            for text in highlighted:
                chunks.append({
                    "content": text,
                    "metadata": {
                        "chunk_type": "highlighted_text",
                        "style": "highlighted"
                    }
                })
            
            # Process strikethrough text
            strikethrough = document_data.get("strikethrough", [])
            for text in strikethrough:
                chunks.append({
                    "content": text,
                    "metadata": {
                        "chunk_type": "strikethrough_text",
                        "style": "strikethrough"
                    }
                })
            
            # Process paragraphs with formatting
            paragraphs = document_data.get("paragraphs", [])
            for para in paragraphs:
                if isinstance(para, dict) and para.get("runs"):
                    has_formatting = any(
                        run.get("bold") or run.get("italic") or 
                        run.get("underline") or run.get("strikethrough") or 
                        run.get("highlight") for run in para.get("runs", [])
                    )
                    
                    if has_formatting:
                        chunks.append({
                            "content": para["text"],
                            "metadata": {
                                "chunk_type": "formatted_text",
                                "formatting_types": [
                                    k for k in ["bold", "italic", "underline", "strikethrough", "highlight"]
                                    if any(run.get(k) for run in para.get("runs", []))
                                ]
                            }
                        })
        
        return chunks

    def store_in_milvus(self, chunks, collection_name="document_chunks"):
        """
        Store document chunks in Milvus vector database.
        """
        docs = []
        for chunk in chunks:
            # Convert chunk to LangChain Document format
            if "title" in chunk:
                text = f"{chunk['title']}\n\n{chunk['content']}"
            else:
                text = chunk['content']
                
            doc = LCDocument(
                page_content=text,
                metadata=chunk.get('metadata', {})
            )
            docs.append(doc)
        
        # Store in Milvus
        vectorstore = Milvus.from_documents(
            docs, 
            self.embeddings, 
            collection_name=collection_name
        )
        return vectorstore
    
    def store_triplets_in_neo4j(self, triplets):
        """
        Store table triplets in Neo4j graph database.
        """
        if not self.neo4j_graph:
            raise ValueError("Neo4j connection not established. Call connect_to_databases first.")
            
        # Clear existing data
        self.neo4j_graph.query("MATCH (n) DETACH DELETE n")
        
        # Create constraints and indices
        self.neo4j_graph.query("CREATE CONSTRAINT IF NOT EXISTS FOR (n:Entity) REQUIRE n.id IS UNIQUE")
        
        # Store each triplet
        for triplet in triplets:
            # Create or merge source and target nodes
            self.neo4j_graph.query(
                """
                MERGE (s:Entity {id: $source})
                MERGE (t:Entity {id: $target})
                CREATE (s)-[:RELATES {type: $relation}]->(t)
                """,
                {"source": triplet["source"], "target": triplet["target"], "relation": triplet["relation"]}
            )
        
        return True
    
    def process_document(self, file_path):
        """
        Process a document file and store its content in vector databases
        according to the specified chunking strategies and storage targets.
        """
        # Step 1: Analyze the document
        document_data = self.analyze_document(file_path)
        
        # Step 2: Apply different chunking strategies
        # Hierarchical chunking for rule explanations (Title + Content)
        hierarchical_chunks = self.hierarchical_chunking(document_data)
        
        # Paragraph chunking for business logic flows and user scenarios
        paragraph_chunks = self.paragraph_chunking(document_data)
        
        # Table chunking for tables (to Neo4j)
        table_triplets = self.table_chunking(document_data)
        
        # Metadata-aware chunking for preserving document structure
        metadata_chunks = self.metadata_aware_chunking(document_data)
        
        # Step 3 & 4: Store chunks in the appropriate vector databases
        # Store rule explanations in Milvus
        rule_chunks = [chunk for chunk in hierarchical_chunks 
                      if chunk.get("metadata", {}).get("chunk_type") == "title_content"]
        rule_vectorstore = self.store_in_milvus(rule_chunks, collection_name="rule_explanations")
        
        # Store business logic flows in Milvus
        business_logic_chunks = [chunk for chunk in paragraph_chunks 
                               if "business" in chunk.get("content", "").lower()]
        business_vectorstore = self.store_in_milvus(business_logic_chunks, collection_name="business_logic_flows")
        
        # Store user scenarios in Milvus
        user_scenario_chunks = [chunk for chunk in paragraph_chunks 
                              if "user" in chunk.get("content", "").lower()]
        user_vectorstore = self.store_in_milvus(user_scenario_chunks, collection_name="user_scenarios")
        
        # Store system interactions in both Neo4j and Milvus
        system_chunks = [chunk for chunk in paragraph_chunks 
                       if "system" in chunk.get("content", "").lower()]
        system_vectorstore = self.store_in_milvus(system_chunks, collection_name="system_interactions")
        
        # For system interactions we could also create graph representations
        system_triplets = []
        for chunk in system_chunks:
            # Simple heuristic to extract potential relationships
            content = chunk.get("content", "")
            matches = re.findall(r"(\w+)\s+(interacts with|calls|uses|depends on|connects to)\s+(\w+)", content)
            for match in matches:
                system_triplets.append({
                    "source": match[0],
                    "relation": match[1],
                    "target": match[2]
                })
        
        # Store tables in Neo4j
        self.store_triplets_in_neo4j(table_triplets + system_triplets)
        
        return {
            "rule_vectorstore": rule_vectorstore,
            "business_vectorstore": business_vectorstore,
            "user_vectorstore": user_vectorstore,
            "system_vectorstore": system_vectorstore
        }


