# Import core libraries
import os
import re
import json
import numpy as np
import pandas as pd
from typing import List, Dict, Tuple, Any, Optional, Union

# LangChain imports
from langchain_core.documents import Document
from langchain_text_splitters import (
    RecursiveCharacterTextSplitter,
    MarkdownHeaderTextSplitter,
    HTMLHeaderTextSplitter
)
from langchain_community.vectorstores.milvus import Milvus
from langchain_neo4j import Neo4jGraph, Neo4jVector
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.retrievers.multi_query import MultiQueryRetriever
from langchain.retrievers import ParentDocumentRetriever
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_google_genai import ChatGoogleGenerativeAI

# Document processing
from unstructured.partition.html import partition_html
from unstructured.partition.pdf import partition_pdf
from bs4 import BeautifulSoup
import PyPDF2
import tabula

# Google Gemini
import google.generativeai as genai
from tqdm.auto import tqdm as notebook_tqdm
print("All required libraries imported successfully.")


# ## 2. Document Analysis and Preprocessing
# We need to create functions to load and analyze document structure, identifying various elements like headings, paragraphs, tables, and images.

# In[14]:


class DocumentAnalyzer:
    """Class to analyze document structure and extract various elements."""

    def __init__(self):
        pass

    def load_document(self, file_path: str) -> Dict:
        """
        Load document from file path and return structured content.

        Args:
            file_path: Path to the document file

        Returns:
            Dictionary with structured document content
        """
        # Get file extension
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()

        if ext == '.pdf':
            return self._process_pdf(file_path)
        elif ext in ['.html', '.htm']:
            return self._process_html(file_path)
        elif ext in ['.txt', '.md']:
            return self._process_text(file_path)
        elif ext in ['.docx', '.doc']:
            return self._process_word(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def _process_pdf(self, file_path: str) -> Dict:
        """Process PDF documents."""
        document_structure = {
            'metadata': {'source': file_path, 'type': 'pdf'},
            'elements': []
        }

        # Read PDF and extract text
        elements = partition_pdf(
            filename=file_path,
            extract_images=True,
            infer_table_structure=True
        )

        # Extract tables separately using tabula
        tables = tabula.read_pdf(file_path, pages='all')

        # Process elements and categorize them
        for element in elements:
            elem_type = str(type(element)).lower()
            element_data = {
                'content': str(element),
                'type': None
            }

            if 'title' in elem_type or 'heading' in elem_type:
                element_data['type'] = 'heading'
            elif 'table' in elem_type:
                element_data['type'] = 'table'
            elif 'image' in elem_type:
                element_data['type'] = 'image' 
            elif 'text' in elem_type:
                # Check if it's strikeout or highlighted (would need more PDF-specific analysis)
                if '~~' in str(element) or '--' in str(element):
                    element_data['type'] = 'strikeout'
                elif any(marker in str(element) for marker in ['**', '__', '>>']):
                    element_data['type'] = 'highlight'
                else:
                    element_data['type'] = 'paragraph'

            document_structure['elements'].append(element_data)

        # Add tables from tabula to our elements list
        for i, table in enumerate(tables):
            document_structure['elements'].append({
                'content': table,
                'type': 'table',
                'pandas_table': True,
                'table_id': i
            })

        return document_structure

    def _process_html(self, file_path: str) -> Dict:
        """Process HTML documents."""
        with open(file_path, 'r', encoding='utf-8') as file:
            html_content = file.read()

        soup = BeautifulSoup(html_content, 'lxml')
        document_structure = {
            'metadata': {'source': file_path, 'type': 'html'},
            'elements': []
        }

        # Extract headings
        for heading_level in range(1, 7):
            for heading in soup.find_all(f'h{heading_level}'):
                document_structure['elements'].append({
                    'content': heading.get_text(),
                    'type': 'heading',
                    'level': heading_level
                })

        # Extract paragraphs
        for para in soup.find_all('p'):
            # Check for highlighted text
            highlighted = para.find_all(['strong', 'b', 'mark', 'em'])
            strikeout = para.find_all('s')

            if highlighted:
                for h in highlighted:
                    document_structure['elements'].append({
                        'content': h.get_text(),
                        'type': 'highlight',
                    })

            if strikeout:
                for s in strikeout:
                    document_structure['elements'].append({
                        'content': s.get_text(),
                        'type': 'strikeout',
                    })

            # Add the full paragraph too
            document_structure['elements'].append({
                'content': para.get_text(),
                'type': 'paragraph',
            })

        # Extract tables
        for table in soup.find_all('table'):
            # Convert HTML table to pandas DataFrame
            table_data = []
            rows = table.find_all('tr')
            for row in rows:
                cols = row.find_all(['td', 'th'])
                cols = [ele.get_text().strip() for ele in cols]
                table_data.append(cols)

            if table_data:
                # Try to create a pandas DataFrame
                try:
                    df = pd.DataFrame(table_data[1:], columns=table_data[0])
                    document_structure['elements'].append({
                        'content': df,
                        'type': 'table',
                        'pandas_table': True
                    })
                except:
                    # Fallback to string representation
                    document_structure['elements'].append({
                        'content': str(table_data),
                        'type': 'table',
                        'pandas_table': False
                    })

        # Extract images
        for img in soup.find_all('img'):
            document_structure['elements'].append({
                'content': img.get('alt', 'Image') + f" (src: {img.get('src', '')})",
                'type': 'image',
            })

        return document_structure

    def _process_text(self, file_path: str) -> Dict:
        """Process plain text or markdown documents."""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()

        document_structure = {
            'metadata': {'source': file_path, 'type': 'text'},
            'elements': []
        }

        # Split by double newlines to separate paragraphs
        paragraphs = content.split('\n\n')

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # Check if it's a heading (starts with # in markdown)
            if para.startswith('#'):
                level = len(re.match(r'^#+', para).group())
                heading_text = para.lstrip('#').strip()
                document_structure['elements'].append({
                    'content': heading_text,
                    'type': 'heading',
                    'level': level
                })
            # Check if it's a table (simple detection for markdown tables)
            elif '|' in para and '-+-' in para.replace('|', '+'):
                document_structure['elements'].append({
                    'content': para,
                    'type': 'table',
                    'pandas_table': False
                })
            # Check for strikeout text (~~text~~ in markdown)
            elif '~~' in para:
                document_structure['elements'].append({
                    'content': para,
                    'type': 'strikeout',
                })
            # Check for highlighted text (** or __ in markdown)
            elif '**' in para or '__' in para:
                document_structure['elements'].append({
                    'content': para,
                    'type': 'highlight',
                })
            # Regular paragraph
            else:
                document_structure['elements'].append({
                    'content': para,
                    'type': 'paragraph',
                })

        return document_structure

    def _process_word(self, file_path: str) -> Dict:
        """
        Process Word documents using python-docx library.
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            Dictionary with structured document content
        """
        try:
            import docx
            from docx.oxml.table import CT_Tbl
            from docx.oxml.text.paragraph import CT_P
            
            document = docx.Document(file_path)
            document_structure = {
                'metadata': {'source': file_path, 'type': 'docx'},
                'elements': []
            }
            
            # Process document body
            for element in document.element.body:
                # Process paragraphs
                if isinstance(element, CT_P):
                    paragraph = docx.Document().add_paragraph()
                    paragraph._p = element
                    text = paragraph.text.strip()
                    
                    if not text:
                        continue
                    
                    # Check if it's likely a heading (based on style)
                    p = document.add_paragraph()
                    p._p = element
                    if hasattr(p, 'style') and p.style and 'heading' in p.style.name.lower():
                        document_structure['elements'].append({
                            'content': text,
                            'type': 'heading',
                            'level': int(p.style.name[-1]) if p.style.name[-1].isdigit() else 1
                        })
                    # Check for highlighted or strikeout text
                    elif '**' in text or '__' in text:
                        document_structure['elements'].append({
                            'content': text,
                            'type': 'highlight'
                        })
                    elif '~~' in text:
                        document_structure['elements'].append({
                            'content': text,
                            'type': 'strikeout'
                        })
                    else:
                        document_structure['elements'].append({
                            'content': text,
                            'type': 'paragraph'
                        })
                
                # Process tables
                elif isinstance(element, CT_Tbl):
                    table = docx.Document().add_table(rows=1, cols=1)
                    table._tbl = element
                    
                    # Convert table to pandas DataFrame
                    data = []
                    headers = []
                    
                    # Get headers from first row
                    if table.rows:
                        for cell in table.rows[0].cells:
                            headers.append(cell.text.strip())
                    
                    # Get data from remaining rows
                    for row in table.rows[1:]:
                        row_data = []
                        for cell in row.cells:
                            row_data.append(cell.text.strip())
                        data.append(row_data)
                    
                    # Create pandas DataFrame if possible
                    try:
                        if headers and data:
                            df = pd.DataFrame(data, columns=headers)
                            document_structure['elements'].append({
                                'content': df,
                                'type': 'table',
                                'pandas_table': True
                            })
                        else:
                            # Create simple text representation for table
                            table_text = "Table content:\n"
                            for row in table.rows:
                                row_text = [cell.text.strip() for cell in row.cells]
                                table_text += " | ".join(row_text) + "\n"
                            document_structure['elements'].append({
                                'content': table_text,
                                'type': 'table',
                                'pandas_table': False
                            })
                    except Exception as e:
                        print(f"Error converting table to DataFrame: {e}")
                        table_text = "Table content (error converting):\n"
                        for row in table.rows:
                            row_text = [cell.text.strip() for cell in row.cells]
                            table_text += " | ".join(row_text) + "\n"
                        document_structure['elements'].append({
                            'content': table_text,
                            'type': 'table',
                            'pandas_table': False
                        })
            
            # Process document properties (metadata)
            try:
                core_properties = document.core_properties
                document_structure['metadata']['title'] = core_properties.title
                document_structure['metadata']['author'] = core_properties.author
                document_structure['metadata']['created'] = str(core_properties.created)
                document_structure['metadata']['modified'] = str(core_properties.modified)
            except Exception as e:
                print(f"Error extracting document properties: {e}")
            
            return document_structure
        
        except ImportError:
            print("python-docx package not found. Please install it with 'pip install python-docx'")
            return {
                'metadata': {'source': file_path, 'type': 'docx'},
                'elements': [{
                    'content': f"python-docx package required for Word processing: {file_path}",
                    'type': 'paragraph',
                }]
            }
        except Exception as e:
            print(f"Error processing Word document: {e}")
            return {
                'metadata': {'source': file_path, 'type': 'docx'},
                'elements': [{
                    'content': f"Error processing Word document: {file_path}. Error: {str(e)}",
                    'type': 'paragraph',
                }]
            }

    def extract_tables_as_triplets(self, document_structure: Dict) -> List[Dict]:
        """
        Extract tables from document and convert to triplets for Neo4j.
        """
        triplets = []
        table_count = 0
        
        try:
            for idx, element in enumerate(document_structure['elements']):
                if element['type'] == 'table':
                    if element.get('pandas_table', False):
                        try:
                            df = element['content']
                            if isinstance(df, pd.DataFrame):
                                # Get table name
                                table_name = f"Table_{table_count}"
                                table_count += 1
                                
                                # Look for the closest heading before this table
                                for j in range(idx-1, -1, -1):
                                    if j < len(document_structure['elements']) and document_structure['elements'][j]['type'] == 'heading':
                                        table_name = document_structure['elements'][j]['content']
                                        break
                                        
                                # Create triplets from table
                                row_count = 0
                                for _, row in df.iterrows():
                                    for col in df.columns:
                                        try:
                                            triplets.append({
                                                'subject': f"{table_name}",
                                                'predicate': str(col),
                                                'object': str(row[col]),
                                                'row_id': row_count
                                            })
                                        except Exception as e:
                                            print(f"Error creating triplet for cell: {e}")
                                    row_count += 1
                        except Exception as e:
                            print(f"Error processing table at index {idx}: {e}")
        except Exception as e:
            print(f"Error in extract_tables_as_triplets: {e}")
            
        return triplets


# ## 3. Implementing Chunking Strategies
# Now we'll develop functions for various chunking strategies:
# 1. Title + Content (Hierarchical) chunking
# 2. Paragraph-Based chunking
# 3. Table-Aware chunking for Neo4j Graph representation
# 4. Metadata-Aware chunking

# In[15]:


class DocumentChunker:
    """Class implementing various chunking strategies for document processing."""

    def __init__(self):
        # Initialize text splitters for different strategies
        self.recursive_splitter = RecursiveCharacterTextSplitter(
            chunk_size=400,
            chunk_overlap=100,
            separators=["\n\n", "\n", ". ", " ", ""]
        )

        self.md_header_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=[
                ("#", "header1"),
                ("##", "header2"),
                ("###", "header3"),
                ("####", "header4")
            ]
        )

        self.html_header_splitter = HTMLHeaderTextSplitter(
            headers_to_split_on=[
                ("h1", "header1"),
                ("h2", "header2"),
                ("h3", "header3"),
                ("h4", "header4")
            ]
        )

    def hierarchical_chunking(self, document_structure: Dict) -> List[Document]:
        """
        Perform Title + Content (Hierarchical) chunking.

        Args:
            document_structure: Document structure dictionary

        Returns:
            List of LangChain Document objects
        """
        # First convert the document structure to a format suitable for hierarchical processing
        markdown_content = self._convert_to_markdown_with_headers(document_structure)

        # Split using the markdown header splitter
        md_header_splits = self.md_header_splitter.split_text(markdown_content)

        # Further split long content chunks if necessary
        final_chunks = []
        for doc in md_header_splits:
            # If content is too long, split further
            content = doc.page_content
            if len(content) > 1000:  # adjust threshold as needed
                sub_chunks = self.recursive_splitter.split_text(content)
                # Copy metadata to each sub-chunk
                for chunk in sub_chunks:
                    final_chunks.append(Document(
                        page_content=chunk,
                        metadata={**doc.metadata, 'chunk_type': 'hierarchical'}
                    ))
            else:
                doc.metadata['chunk_type'] = 'hierarchical'
                final_chunks.append(doc)

        return final_chunks

    def paragraph_based_chunking(self, document_structure: Dict) -> List[Document]:
        """
        Perform Paragraph-Based chunking with line breaks preserved.

        Args:
            document_structure: Document structure dictionary

        Returns:
            List of LangChain Document objects
        """
        chunks = []

        # Extract paragraphs from the document
        paragraphs = []
        current_section = {'title': 'Document', 'content': ''}

        for element in document_structure['elements']:
            if element['type'] == 'heading':
                # If we have content in the current section, save it
                if current_section['content'].strip():
                    paragraphs.append(current_section)

                # Start a new section
                current_section = {
                    'title': element['content'],
                    'content': ''
                }
            elif element['type'] == 'paragraph':
                current_section['content'] += element['content'] + "\n\n"

        # Add the last section if it has content
        if current_section['content'].strip():
            paragraphs.append(current_section)

        # Create Document objects for each paragraph
        for para in paragraphs:
            # Split content if it's too long
            if len(para['content']) > 1000:
                sub_chunks = self.recursive_splitter.split_text(para['content'])
                for i, chunk in enumerate(sub_chunks):
                    chunks.append(Document(
                        page_content=chunk,
                        metadata={
                            'title': para['title'],
                            'chunk_type': 'paragraph',
                            'chunk_index': i,
                            'source': document_structure['metadata']['source']
                        }
                    ))
            else:
                chunks.append(Document(
                    page_content=para['content'],
                    metadata={
                        'title': para['title'],
                        'chunk_type': 'paragraph',
                        'source': document_structure['metadata']['source']
                    }
                ))

        return chunks

    def table_aware_chunking(self, document_structure: Dict) -> List[Dict]:
        """
        Perform Table-Aware chunking for Neo4j graph.

        Args:
            document_structure: Document structure dictionary

        Returns:
            List of dictionaries representing table triplets for Neo4j
        """
        # Extract tables and convert to triplets
        analyzer = DocumentAnalyzer()
        return analyzer.extract_tables_as_triplets(document_structure)

    def metadata_aware_chunking(self, document_structure: Dict) -> List[Document]:
        """
        Perform Metadata-Aware chunking with Milvus compatibility.
        """
        chunks = []
        
        # Identify semantic sections
        semantic_sections = self._identify_semantic_sections(document_structure)
        
        for section in semantic_sections:
            # Calculate chunk size based on content type
            if section['type'] == 'rule_explanation':
                chunk_size = 1500
            elif section['type'] == 'user_scenario':
                chunk_size = 1000
            elif section['type'] == 'business_logic':
                chunk_size = 800
            else:
                chunk_size = 1000
                
            custom_splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=min(100, chunk_size // 10),
                separators=["\n\n", "\n", ". ", " ", ""]
            )
            
            # Split content
            text_chunks = custom_splitter.split_text(section['content'])
            
            # Convert keywords and categories to strings for Milvus compatibility
            keywords_str = ",".join(section.get('keywords', []))
            categories_str = ",".join(section.get('categories', []))
            
            # Create Document objects with compatible metadata
            for i, chunk in enumerate(text_chunks):
                chunks.append(Document(
                    page_content=chunk,
                    metadata={
                        'title': section.get('title', 'Untitled Section'),
                        'content_type': section['type'],
                        'source': document_structure['metadata']['source'],
                        'chunk_index': i,
                        'total_chunks': len(text_chunks),
                        'chunk_type': 'metadata_aware',
                        # Store keywords and categories as strings
                        'keywords_str': keywords_str,
                        'categories_str': categories_str
                    }
                ))
        
        return chunks

    def _convert_to_markdown_with_headers(self, document_structure: Dict) -> str:
        """
        Convert document structure to markdown format with headers.

        Args:
            document_structure: Document structure dictionary

        Returns:
            Markdown string representation of the document
        """
        markdown = []
        current_heading = ""

        for element in document_structure['elements']:
            if element['type'] == 'heading':
                level = element.get('level', 1)
                heading = '#' * level + ' ' + element['content']
                markdown.append(heading)
                current_heading = element['content']
            elif element['type'] == 'paragraph':
                markdown.append(element['content'])
            elif element['type'] == 'table':
                if isinstance(element['content'], pd.DataFrame):
                    table_str = element['content'].to_markdown()
                    markdown.append(table_str)
                else:
                    markdown.append(str(element['content']))
            elif element['type'] in ['strikeout', 'highlight', 'image']:
                markdown.append(element['content'])

        return '\n\n'.join(markdown)

    def _identify_semantic_sections(self, document_structure: Dict) -> List[Dict]:
        """
        Identify semantic sections in the document.

        Args:
            document_structure: Document structure dictionary

        Returns:
            List of semantic section dictionaries
        """
        sections = []
        current_section = None

        # Patterns to identify content types
        rule_patterns = ['rule', 'regulation', 'policy', 'guidelines']
        logic_patterns = ['workflow', 'process', 'logic', 'procedure', 'business logic']
        scenario_patterns = ['scenario', 'example', 'use case', 'user story']
        interaction_patterns = ['interaction', 'interface', 'system', 'integration', 'api']

        for element in document_structure['elements']:
            if element['type'] == 'heading':
                # If there's a current section with content, save it
                if current_section and current_section.get('content'):
                    sections.append(current_section)

                # Determine the section type based on heading content
                heading_lower = element['content'].lower()

                if any(pattern in heading_lower for pattern in rule_patterns):
                    section_type = 'rule_explanation'
                elif any(pattern in heading_lower for pattern in logic_patterns):
                    section_type = 'business_logic'
                elif any(pattern in heading_lower for pattern in scenario_patterns):
                    section_type = 'user_scenario'
                elif any(pattern in heading_lower for pattern in interaction_patterns):
                    section_type = 'system_interaction'
                else:
                    section_type = 'general'

                # Create a new section
                current_section = {
                    'title': element['content'],
                    'type': section_type,
                    'content': '',
                    'keywords': self._extract_keywords(element['content']),
                    'categories': [section_type]
                }
            elif element['type'] in ['paragraph', 'strikeout', 'highlight']:
                if current_section:
                    current_section['content'] += element['content'] + "\n\n"
                else:
                    # Create a default section if none exists
                    current_section = {
                        'title': 'Introduction',
                        'type': 'general',
                        'content': element['content'] + "\n\n",
                        'keywords': [],
                        'categories': ['general']
                    }
            elif element['type'] == 'table':
                # Tables are handled separately for Neo4j, but we still include their text representation
                # in the content for completeness
                if current_section:
                    if isinstance(element['content'], pd.DataFrame):
                        current_section['content'] += "Table content:\n"
                        current_section['content'] += str(element['content']) + "\n\n"
                    else:
                        current_section['content'] += "Table content:\n" + str(element['content']) + "\n\n"

        # Add the last section if it exists
        if current_section and current_section.get('content'):
            sections.append(current_section)

        return sections

    def _extract_keywords(self, text: str) -> List[str]:
        """
        Extract keywords from text.

        Args:
            text: Text to extract keywords from

        Returns:
            List of keywords
        """
        # Simple implementation - extract meaningful words and filter stop words
        stop_words = {'a', 'an', 'the', 'and', 'or', 'but', 'if', 'then', 'else', 'when', 
                     'on', 'in', 'at', 'to', 'for', 'with', 'by', 'about', 'as', 'of'}

        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
        keywords = [word for word in words if word not in stop_words]

        # Return unique keywords
        return list(set(keywords))


# ## 4. Setting Up the Embedding Model
# We'll configure a local embedding model to generate vector representations of our chunks.

# In[16]:


class EmbeddingManager:
    """Class to manage embeddings for document chunks."""

    def __init__(self, model_name="all-MiniLM-L6-v2"):
        """
        Initialize the embedding manager.

        Args:
            model_name: Name of the HuggingFace model to use for embeddings
        """
        self.model_name = model_name
        self.embeddings = HuggingFaceEmbeddings(model_name=model_name)
        print(f"Initialized embedding model: {model_name}")

    def generate_embedding(self, text: str) -> List[float]:
        """
        Generate an embedding for text.

        Args:
            text: Text to embed

        Returns:
            List of floats representing the embedding vector
        """
        # The HuggingFaceEmbeddings class follows the LangChain embedding interface
        return self.embeddings.embed_query(text)

    def batch_embed(self, texts: List[str]) -> List[List[float]]:
        """
        Batch embed multiple texts.

        Args:
            texts: List of texts to embed

        Returns:
            List of embedding vectors
        """
        return self.embeddings.embed_documents(texts)

    def embed_documents(self, documents: List[Document]) -> Tuple[List[List[float]], List[Document]]:
        """
        Embed LangChain Document objects.

        Args:
            documents: List of Document objects to embed

        Returns:
            Tuple of (embeddings, documents)
        """
        texts = [doc.page_content for doc in documents]
        embeddings = self.batch_embed(texts)
        return embeddings, documents

    def embed_triplets(self, triplets: List[Dict]) -> List[Dict]:
        """
        Embed triplets for Neo4j graph.

        Args:
            triplets: List of triplet dictionaries

        Returns:
            Triplets with embeddings added
        """
        for triplet in triplets:
            # Create a concatenated text representation of the triplet
            triplet_text = f"{triplet['subject']} {triplet['predicate']} {triplet['object']}"
            triplet['embedding'] = self.generate_embedding(triplet_text)

        return triplets


# ## 5. Configuring Vector Databases (Milvus and Neo4j)
# Let's set up connections and schemas for Milvus and Neo4j databases.

# In[17]:


class VectorDBManager:
    """Class to manage vector database connections and operations."""

    def __init__(self, embedding_manager: EmbeddingManager):
        """
        Initialize the vector database manager.

        Args:
            embedding_manager: Embedding manager to use for generating embeddings
        """
        self.embedding_manager = embedding_manager
        self.milvus_client = None
        self.neo4j_client = None

        # Collections in Milvus
        self.collections = {
            'rule_explanation': None,
            'business_logic': None,
            'user_scenario': None,
            'system_interaction': None
        }

    def setup_milvus(self, host='localhost', port='2379'):
        """
        Set up connection to Milvus.
        
        Args:
            host: Milvus host
            port: Milvus port
        """
        try:
            # Define Milvus connection parameters
            self.milvus_connection_params = {
                "host": host,
                "port": port
            }
            
            print("Milvus connection parameters configured. Collections will be created when data is inserted.")
            self.milvus_available = True
        except Exception as e:
            print(f"Error setting up Milvus: {e}")
            self.milvus_available = False

    def setup_neo4j(self, uri='bolt://localhost:7687', username='neo4j', password='venev'):
        """
        Set up connection to Neo4j.
        
        Args:
            uri: Neo4j URI
            username: Neo4j username
            password: Neo4j password
        """
        try:
            # Store Neo4j connection parameters
            self.neo4j_connection_params = {
                "url": uri,
                "username": username,
                "password": password
            }
            
            # Initialize Neo4j graph client from LangChain
            self.neo4j_client = Neo4jGraph(
                url=uri,
                username=username,
                password=password
            )
            
            # Initialize Neo4j vector client with correct parameters
            self.neo4j_vector = Neo4jVector(
                url=uri,
                username=username,
                password=password,
                embedding=self.embedding_manager.embeddings,
                index_name="documentVectors",
                node_label="DocumentChunk"
            )
            
            # Set up schema for our document graph
            self._create_neo4j_schema()
            print("Neo4j connection established and schema created.")
            self.neo4j_available = True
        except Exception as e:
            print(f"Error setting up Neo4j: {e}")
            self.neo4j_client = None
            self.neo4j_vector = None
            self.neo4j_available = False

    def _create_neo4j_schema(self):
        """Create Neo4j schema for document storage."""
        # Create constraints that work with Community Edition
        self.neo4j_client.query("""
            CREATE CONSTRAINT document_id IF NOT EXISTS
            FOR (d:Document) REQUIRE d.id IS UNIQUE
        """)
        
        self.neo4j_client.query("""
            CREATE CONSTRAINT chunk_id IF NOT EXISTS
            FOR (c:DocumentChunk) REQUIRE c.id IS UNIQUE
        """)
        
        # Remove the existence constraint that requires Enterprise Edition
        # self.neo4j_client.query("""
        #    CREATE CONSTRAINT documentchunk_text IF NOT EXISTS
        #    FOR (c:DocumentChunk) REQUIRE c.text IS NOT NULL
        # """)
        
        # Create vector index if needed
        self.neo4j_client.query("""
            CREATE VECTOR INDEX documentVectors IF NOT EXISTS
            FOR (c:DocumentChunk) 
            ON c.embedding
            OPTIONS {indexConfig: {
              `vector.dimensions`: 384,
              `vector.similarity_function`: 'cosine'
            }}
        """)

    def store_in_milvus(self, documents: List[Document], content_type: str):
        """
        Store documents in Milvus.
        
        Args:
            documents: List of Document objects to store
            content_type: Type of content (rule_explanation, business_logic, etc.)
        """
        if not hasattr(self, 'milvus_available') or not self.milvus_available:
            print(f"Skipping Milvus storage for {content_type} - Milvus not available")
            return
            
        try:
            # Check if content type is valid
            if content_type not in self.collections:
                raise ValueError(f"Invalid content type: {content_type}")

            # Create collection if it doesn't exist
            collection_name = f"docs_{content_type}"

            # Generate embeddings for documents
            embeddings = [self.embedding_manager.generate_embedding(doc.page_content) for doc in documents]

            # Create Milvus collection
            vector_store = Milvus.from_documents(
                documents=documents,
                embedding=self.embedding_manager.embeddings,
                collection_name=collection_name,
                connection_args=self.milvus_connection_params
            )

            # Store the collection reference
            self.collections[content_type] = vector_store

            print(f"Stored {len(documents)} documents in Milvus collection: {collection_name}")

        except Exception as e:
            print(f"Error storing documents in Milvus: {e}")

    def store_in_neo4j(self, triplets: List[Dict]):
        """
        Store triplets in Neo4j.
        
        Args:
            triplets: List of triplet dictionaries to store
        """
        if not hasattr(self, 'neo4j_available') or not self.neo4j_available:
            print("Skipping Neo4j storage - Neo4j not available")
            return
            
        try:
            # Get the Neo4j driver session properly
            # The neo4j_client has a session() method on its driver
            session = self.neo4j_client._driver.session()
            
            # Process in smaller batches for better error handling
            batch_size = 50
            for i in range(0, len(triplets), batch_size):
                batch = triplets[i:i+batch_size]
                try:
                    # Begin a transaction for this batch
                    with session.begin_transaction() as tx:
                        # Convert parameters for Neo4j
                        params = {
                            "batch_size": len(batch),
                            "subjects": [t["subject"] for t in batch],
                            "predicates": [t["predicate"] for t in batch],
                            "objects": [t["object"] for t in batch],
                            "embeddings": [t["embedding"] for t in batch]
                        }

                        # Fixed Cypher query
                        query = """
                        UNWIND range(0, $batch_size - 1) as i
                        MERGE (s:Subject {name: $subjects[i]})
                        MERGE (o:Object {name: $objects[i]})
                        WITH s, o, i
                        CREATE (s)-[r:HAS_PROPERTY {name: $predicates[i]}]->(o)
                        WITH s, o, i
                        CREATE (c:DocumentChunk {
                            id: randomUUID(),
                            text: $subjects[i] + ' ' + $predicates[i] + ' ' + $objects[i],
                            subject: $subjects[i],
                            predicate: $predicates[i],
                            object: $objects[i]
                        })
                        SET c.embedding = $embeddings[i]
                        CREATE (c)-[:SUBJECT_OF]->(s)
                        CREATE (c)-[:OBJECT_OF]->(o)
                        """

                        # Execute query within transaction
                        result = tx.run(query, params)
                        result.consume()  # Ensure execution completes
                    
                    print(f"Processed batch {i//batch_size + 1}/{(len(triplets) + batch_size - 1)//batch_size}")
                        
                except Exception as e:
                    print(f"Error processing batch {i//batch_size + 1}: {e}")
                    # Continue to the next batch even if this one fails
            
            # Close the session when done
            session.close()
            print(f"Stored triplets in Neo4j (completed batches)")

        except Exception as e:
            print(f"Error initializing Neo4j transaction: {e}")
            # Make sure session is closed if an error occurs
            if 'session' in locals():
                session.close()

    def hybrid_retrieval(self, query: str, content_types: List[str], k: int = 5):
        """
        Perform hybrid retrieval from Milvus and Neo4j.

        Args:
            query: Query string
            content_types: List of content types to search
            k: Number of results to return from each source

        Returns:
            List of retrieved documents
        """
        results = []

        # Retrieve from Milvus
        for content_type in content_types:
            if content_type in self.collections and self.collections[content_type] is not None:
                milvus_results = self.collections[content_type].similarity_search(query, k=k)
                results.extend(milvus_results)

        # Retrieve from Neo4j if available
        if self.neo4j_vector:
            neo4j_results = self.neo4j_vector.similarity_search(query, k=k)
            results.extend(neo4j_results)

        return results


# ## 6. Document Processing Pipeline
# Now let's develop a pipeline that processes documents by content type, applies appropriate chunking strategies, and stores data in the correct database.

# In[18]:


class DocumentProcessingPipeline:
    """Pipeline for document processing, chunking, and storage."""

    def __init__(
        self, 
        analyzer: DocumentAnalyzer, 
        chunker: DocumentChunker, 
        embedding_manager: EmbeddingManager,
        vector_db_manager: VectorDBManager
    ):
        """
        Initialize the document processing pipeline.

        Args:
            analyzer: Document analyzer
            chunker: Document chunker
            embedding_manager: Embedding manager
            vector_db_manager: Vector database manager
        """
        self.analyzer = analyzer
        self.chunker = chunker
        self.embedding_manager = embedding_manager
        self.vector_db_manager = vector_db_manager

        # Content type mapping
        self.content_mapping = {
            'rule_explanation': {'chunking': 'hierarchical', 'storage': 'milvus'},
            'business_logic': {'chunking': 'paragraph', 'storage': 'milvus'},
            'table': {'chunking': 'table', 'storage': 'neo4j'},
            'user_scenario': {'chunking': 'paragraph', 'storage': 'milvus'},
            'system_interaction': {'chunking': 'metadata', 'storage': 'both'}
        }

    def process_document(self, file_path: str):
        """
        Process a document file through the entire pipeline.

        Args:
            file_path: Path to the document file
        """
        print(f"Processing document: {file_path}")

        # Step 1: Analyze document
        document_structure = self.analyzer.load_document(file_path)
        print(f"Document analysis complete. Found {len(document_structure['elements'])} elements.")

        # Step 2: Process each content type with appropriate chunking
        content_types = self._identify_content_types(document_structure)
        print(f"Identified content types: {', '.join(content_types)}")

        # Step 3: Apply chunking strategies and store in appropriate databases
        self._process_content_types(document_structure, content_types)

        print("Document processing complete.")

    def _identify_content_types(self, document_structure: Dict) -> List[str]:
        """
        Identify content types in the document structure.

        Args:
            document_structure: Document structure dictionary

        Returns:
            List of identified content types
        """
        content_types = set()

        # Look for tables explicitly
        for element in document_structure['elements']:
            if element['type'] == 'table':
                content_types.add('table')

        # For other content types, we need more sophisticated analysis
        # Convert document to text for analysis
        document_text = ""
        for element in document_structure['elements']:
            if element['type'] in ['heading', 'paragraph']:
                document_text += element['content'] + "\n\n"

        # Check for rule explanations
        if any(keyword in document_text.lower() for keyword in 
               ['rule', 'regulation', 'policy', 'guidelines', 'requirement']):
            content_types.add('rule_explanation')

        # Check for business logic
        if any(keyword in document_text.lower() for keyword in 
               ['workflow', 'process', 'logic', 'procedure', 'if then', 'business logic']):
            content_types.add('business_logic')

        # Check for user scenarios
        if any(keyword in document_text.lower() for keyword in 
               ['scenario', 'example', 'use case', 'user story', 'user journey']):
            content_types.add('user_scenario')

        # Check for system interactions
        if any(keyword in document_text.lower() for keyword in 
               ['interaction', 'interface', 'system', 'integration', 'api', 'endpoint']):
            content_types.add('system_interaction')

        # If no specific types identified, default to business_logic
        if not content_types or (content_types == {'table'}):
            content_types.add('business_logic')

        return list(content_types)

    def _process_content_types(self, document_structure: Dict, content_types: List[str]):
        """Process content types with improved error handling."""
        # Create a copy of content_types to avoid modification during iteration
        types_to_process = content_types.copy()
        
        # Process tables first if present
        if 'table' in types_to_process:
            print("Processing tables...")
            try:
                triplets = self.chunker.table_aware_chunking(document_structure)
                
                if triplets:
                    # Process triplets as before...
                    if hasattr(self.vector_db_manager, 'neo4j_available') and self.vector_db_manager.neo4j_available:
                        try:
                            triplets_with_embeddings = self.embedding_manager.embed_triplets(triplets)
                            self.vector_db_manager.store_in_neo4j(triplets_with_embeddings)
                        except Exception as e:
                            print(f"Error storing triplets in Neo4j: {e}")
                            
                    # Fallback to Milvus regardless, for redundancy
                    try:
                        print("Storing table data in Milvus for backup")
                        table_chunks = []
                        for triplet in triplets:
                            table_chunks.append(Document(
                                page_content=f"{triplet['subject']} - {triplet['predicate']}: {triplet['object']}",
                                metadata={
                                    'title': triplet['subject'],
                                    'chunk_type': 'table_fallback',
                                    'source': document_structure['metadata']['source']
                                }
                            ))
                        if table_chunks:
                            self.vector_db_manager.store_in_milvus(table_chunks, 'business_logic')
                    except Exception as e:
                        print(f"Error storing table fallback in Milvus: {e}")
                else:
                    print("No valid tables found for extraction")
                    
            except Exception as e:
                print(f"Error processing tables: {e}")
                
            types_to_process.remove('table')
        
        # Process other content types with better error handling
        for content_type in types_to_process:
            print(f"Processing content type: {content_type}")
            
            try:
                chunking_strategy = self.content_mapping[content_type]['chunking']
                storage = self.content_mapping[content_type]['storage']
                
                # Apply chunking strategy with error handling
                chunks = []
                try:
                    if chunking_strategy == 'hierarchical':
                        chunks = self.chunker.hierarchical_chunking(document_structure)
                    elif chunking_strategy == 'paragraph':
                        chunks = self.chunker.paragraph_based_chunking(document_structure)
                    elif chunking_strategy == 'metadata':
                        chunks = self.chunker.metadata_aware_chunking(document_structure)
                    else:
                        print(f"Unknown chunking strategy for {content_type}: {chunking_strategy}")
                        continue
                        
                    print(f"Created {len(chunks)} chunks using {chunking_strategy} strategy")
                except Exception as e:
                    print(f"Error applying chunking strategy '{chunking_strategy}': {e}")
                    continue
                    
                # Store in appropriate database with fallbacks
                if not chunks:
                    print(f"No chunks created for {content_type}. Skipping storage.")
                    continue
                    
                try:
                    # Always try Milvus first for reliability
                    if hasattr(self.vector_db_manager, 'milvus_available') and self.vector_db_manager.milvus_available:
                        print(f"Storing {len(chunks)} chunks in Milvus")
                        self.vector_db_manager.store_in_milvus(chunks, content_type)
                        
                    # If Neo4j storage is required and available
                    if (storage == 'neo4j' or storage == 'both'):
                        try:
                            print(f"Converting chunks to triplets for Neo4j storage")
                            triplets = self._convert_chunks_to_triplets(chunks, content_type)
                            triplets_with_embeddings = self.embedding_manager.embed_triplets(triplets)
                            
                            # Try Neo4j storage
                            if hasattr(self.vector_db_manager, 'neo4j_available') and self.vector_db_manager.neo4j_available:
                                try:
                                    self.vector_db_manager.store_in_neo4j(triplets_with_embeddings)
                                except Exception as e:
                                    print(f"Error storing in Neo4j: {e}")
                                    print("Falling back to Milvus storage for this content")
                                    # Always ensure data is stored somewhere
                                    if hasattr(self.vector_db_manager, 'milvus_available') and self.vector_db_manager.milvus_available:
                                        triplet_chunks = []
                                        for triplet in triplets:
                                            triplet_chunks.append(Document(
                                                page_content=f"{triplet['subject']} {triplet['predicate']} {triplet['object']}",
                                                metadata={
                                                    'title': f"{triplet['subject']}",
                                                    'chunk_type': 'neo4j_fallback',
                                                    'source': document_structure['metadata']['source'] 
                                                }
                                            ))
                                        if triplet_chunks:
                                            # Store backup in Milvus
                                            self.vector_db_manager.store_in_milvus(triplet_chunks, content_type)
                                            
                        except Exception as e:
                            print(f"Error preparing triplets: {e}")
                except Exception as e:
                    print(f"Error storing chunks for {content_type}: {e}")
            except Exception as e:
                print(f"Error processing content type {content_type}: {e}")

    def _convert_chunks_to_triplets(self, chunks: List[Document], content_type: str) -> List[Dict]:
        """
        Convert chunks to triplets for Neo4j storage.

        Args:
            chunks: List of Document objects
            content_type: Type of content

        Returns:
            List of triplet dictionaries
        """
        triplets = []

        for chunk in chunks:
            # The approach to convert to triplets depends on the content
            if content_type == 'system_interaction':
                # For system interactions, create subject-action-object triplets
                # This is a simplified approach; in practice, you might use an NLP model
                title = chunk.metadata.get('title', 'System Interaction')
                content = chunk.page_content

                # Simple rule-based extraction of subject-verb-object
                for sentence in content.split('. '):
                    if len(sentence.strip()) < 10:
                        continue

                    # Very simple SVO extraction (would use proper NLP in production)
                    parts = sentence.strip().split(' ', 2)
                    if len(parts) >= 3:
                        subject = parts[0]
                        predicate = parts[1]
                        object_val = parts[2]

                        triplets.append({
                            'subject': f"{title}_{subject}",
                            'predicate': predicate,
                            'object': object_val
                        })
                    else:
                        # Fallback for sentences that don't match our pattern
                        triplets.append({
                            'subject': title,
                            'predicate': 'contains',
                            'object': sentence.strip()
                        })
            else:
                # For other content types, use a simpler approach
                title = chunk.metadata.get('title', 'Document Section')
                content = chunk.page_content

                triplets.append({
                    'subject': title,
                    'predicate': 'contains',
                    'object': content
                })

        return triplets


# ## 7. Hybrid Retrieval System with LangChain
# Now let's implement a hybrid retrieval system using LangChain v0.3 to combine search results from both Milvus and Neo4j.

# In[19]:


class HybridRetriever:
    """Class for hybrid retrieval from multiple vector databases."""

    def __init__(self, vector_db_manager: VectorDBManager):
        """
        Initialize the hybrid retriever.

        Args:
            vector_db_manager: Vector database manager instance
        """
        self.vector_db_manager = vector_db_manager

    def setup_langchain_retrieval(self, use_gemini_for_query_expansion=True):
        """Set up LangChain retrieval with better error handling."""
        try:
            # Create retriever objects for different content types and databases
            self.retrievers = {}

            # Add Milvus retrievers
            for content_type, collection in self.vector_db_manager.collections.items():
                if collection is not None:
                    try:
                        self.retrievers[f"milvus_{content_type}"] = collection.as_retriever(
                            search_type="similarity",
                            search_kwargs={"k": 3}
                        )
                    except Exception as e:
                        print(f"Error setting up retriever for {content_type}: {e}")

            # Add Neo4j retriever if available
            if hasattr(self.vector_db_manager, 'neo4j_vector') and self.vector_db_manager.neo4j_vector:
                try:
                    self.retrievers["neo4j"] = self.vector_db_manager.neo4j_vector.as_retriever(
                        search_kwargs={"k": 3}
                    )
                except Exception as e:
                    print(f"Error setting up Neo4j retriever: {e}")
            
            # Set up query expansion if requested
            if use_gemini_for_query_expansion:
                try:
                    if "GOOGLE_API_KEY" in os.environ and os.environ["GOOGLE_API_KEY"]:
                        # Set up Gemini for multi-query generation
                        self.gemini_llm = ChatGoogleGenerativeAI(
                            model="gemini-1.5-flash-8b",
                            temperature=0.2
                        )

                        # Set up for main retrievers
                        for retriever_name in list(self.retrievers.keys()):
                            if retriever_name.startswith("milvus_"):
                                try:
                                    self.multi_query_retriever = MultiQueryRetriever.from_llm(
                                        retriever=self.retrievers[retriever_name],
                                        llm=self.gemini_llm
                                    )
                                    print("Multi-query retriever set up with Gemini")
                                    break
                                except Exception as e:
                                    print(f"Error setting up multi-query with {retriever_name}: {e}")
                    else:
                        print("No Google API key found. Skipping query expansion setup.")
                except Exception as e:
                    print(f"Error setting up query expansion: {e}")

            print(f"Set up {len(self.retrievers)} retrievers for hybrid search")

        except Exception as e:
            print(f"Error setting up LangChain retrieval: {e}")
            # Initialize empty retrievers to avoid errors
            self.retrievers = {}

    def hybrid_retrieve(self, query: str, k: int = 5):
        """
        Perform hybrid retrieval from all configured retrievers.
        
        Args:
            query: Query string
            k: Number of results to return per retriever
            
        Returns:
            Combined list of unique retrieved documents
        """
        all_results = []
        active_retrievers = 0
        
        # Check if there are any retrievers configured
        if not hasattr(self, 'retrievers') or not self.retrievers:
            print("No retrievers configured. Returning empty results.")
            return []
            
        # Use multi-query retriever if available
        if hasattr(self, 'multi_query_retriever'):
            try:
                print("Using multi-query retriever")
                multi_results = self.multi_query_retriever.invoke(query)
                all_results.extend(multi_results)
                active_retrievers += 1
            except Exception as e:
                print(f"Error with multi-query retrieval: {e}")
        
        # Use all other retrievers
        for name, retriever in self.retrievers.items():
            try:
                print(f"Querying retriever: {name}")
                results = retriever.invoke(query)  # Using invoke instead of get_relevant_documents
                all_results.extend(results)
                active_retrievers += 1
            except Exception as e:
                print(f"Error retrieving from {name}: {e}")
        
        if active_retrievers == 0:
            print("Warning: No retrievers were successfully queried.")
        
        # Remove duplicates (simplified - in practice should use better deduplication)
        seen_contents = set()
        unique_results = []
        
        for doc in all_results:
            content = doc.page_content
            if content not in seen_contents:
                seen_contents.add(content)
                unique_results.append(doc)
        
        # Return the top k results
        return unique_results[:k]

    def create_langchain_rag_pipeline(self):
        """
        Create a LangChain RAG pipeline with hybrid retrieval.

        Returns:
            Runnable chain for RAG
        """
        # Set up Gemini LLM if API key is available
        if "GOOGLE_API_KEY" in os.environ:
            print("Creating LangChain RAG pipeline with Gemini")

            llm = ChatGoogleGenerativeAI(
                model="gemini-1.5-flash-8b",
                temperature=0.7,
                top_p=0.95,
                top_k=40,
                max_output_tokens=2048,
            )

            # Create a prompt template
            template = """
            You are an AI assistant that provides accurate and helpful information based on the given context.

            Context information:
            {context}

            Question: {question}

            If the context doesn't contain relevant information, please say "I don't have enough information to answer this question."
            Provide a comprehensive answer based on the context provided:
            """

            prompt = ChatPromptTemplate.from_template(template)

            # Create a retriever that combines results from all our retrievers
            def retrieve_docs(query):
                return self.hybrid_retrieve(query, k=5)

            # Build the RAG chain
            rag_chain = (
                {"context": retrieve_docs, "question": RunnablePassthrough()}
                | prompt
                | llm
                | StrOutputParser()
            )

            return rag_chain
        else:
            print("Google API key not found, cannot create RAG pipeline")
            return None


# ## 8. Query Processing with Gemini LLM
# Let's integrate the Gemini-1.5-flash-8b model via API to process queries against the retrieved document chunks.

# In[20]:


class GeminiProcessor:
    """Class for processing queries with Gemini LLM."""

    def __init__(self, api_key=None):
        """
        Initialize the Gemini processor.

        Args:
            api_key: Google API key for Gemini
        """
        # Try to get API key from parameter, then environment variable
        self.api_key = api_key or os.environ.get("GOOGLE_API_KEY")
        
        # Remove the streamlit dependency
        if not self.api_key:
            print("Warning: No API key provided for Gemini. Set GOOGLE_API_KEY environment variable.")
        else:
            try:
                genai.configure(api_key=self.api_key)
                self.model = genai.GenerativeModel('gemini-1.5-flash-8b')
                print("Gemini model initialized successfully")
            except Exception as e:
                print(f"Error initializing Gemini: {e}")
                self.model = None

    def process_query(self, query: str, retrieved_docs: List[Document]) -> str:
        """
        Process a query using the Gemini LLM and retrieved documents.

        Args:
            query: User query
            retrieved_docs: List of retrieved Document objects

        Returns:
            Response from Gemini
        """
        if not self.model:
            return "Gemini model not initialized. Please provide a valid API key."

        try:
            # Prepare context from retrieved documents
            context = "\n\n".join([f"Document {i+1}:\n{doc.page_content}" for i, doc in enumerate(retrieved_docs)])

            # Format prompt for Gemini
            prompt = f"""
            You are an intelligent assistant that answers questions based on provided context.

            Context information:
            {context}
            User question: {query}
            Please answer the question based only on the provided context. If the context doesn't contain enough information 
            to fully answer the question, acknowledge what you can answer based on the context and what information is missing.
            """

            # Generate response
            response = self.model.generate_content(prompt)
            return response.text

        except Exception as e:
            print(f"Error processing query with Gemini: {e}")
            return f"Error processing query: {str(e)}"

    def generate_query_variations(self, query: str, num_variations: int = 3) -> List[str]:
        """
        Generate variations of a query for better retrieval.

        Args:
            query: Original query
            num_variations: Number of variations to generate

        Returns:
            List of query variations
        """
        if not self.model:
            return [query]  # Return original query if model not available

        try:
            prompt = f"""
            Generate {num_variations} different ways to ask the following question. 
            Return only the list of questions without any explanations or numbering.
            Original question: {query}
            """
            response = self.model.generate_content(prompt)
            variations = [line.strip() for line in response.text.split('\n') if line.strip()]
            # Include the original query
            if query not in variations:
                variations.insert(0, query)

            # Limit to requested number
            return variations[:num_variations]
        except Exception as e:
            print(f"Error generating query variations: {e}")
            return [query]  # Return original query on error


# ## 9. Demo and Usage Examples
# Let's create a complete demonstration of the system functionality.



def process_document(file_path: str, milvus_host='localhost', milvus_port='19530', 
                    neo4j_uri='bolt://localhost:7687', neo4j_username='neo4j', neo4j_password='venev'):
    """Process a document through the entire pipeline."""
    # Initialize components
    analyzer = DocumentAnalyzer()
    chunker = DocumentChunker()
    embedding_manager = EmbeddingManager(model_name="all-MiniLM-L6-v2")
    
    # Initialize vector database manager
    vector_db_manager = VectorDBManager(embedding_manager)
    
    # Connect to databases with provided connection parameters
    # Add error handling so one failure doesn't stop everything
    try:
        vector_db_manager.setup_milvus(host=milvus_host, port=milvus_port)
    except Exception as e:
        print(f"Warning: Failed to set up Milvus: {e}")
    
    try:
        vector_db_manager.setup_neo4j(uri=neo4j_uri, username=neo4j_username, neo4j_password=neo4j_password)
    except Exception as e:
        print(f"Warning: Failed to set up Neo4j: {e}")
    
    # Create and run the processing pipeline
    pipeline = DocumentProcessingPipeline(
        analyzer=analyzer,
        chunker=chunker, 
        embedding_manager=embedding_manager,
        vector_db_manager=vector_db_manager
    )
    
    # Process the document
    try:
        pipeline.process_document(file_path)
    except Exception as e:
        print(f"Error during document processing: {e}")
    
    # Set up retrieval system
    retriever = HybridRetriever(vector_db_manager)
    
    try:
        retriever.setup_langchain_retrieval(use_gemini_for_query_expansion=True)
    except Exception as e:
        print(f"Error setting up LangChain retrieval: {e}")
    
    return retriever


def query_document(retriever, query: str, k: int = 5):
    """
    Query the document storage system.
    
    Args:
        retriever: HybridRetriever instance from process_document
        query: Query string to search for
        k: Number of results to return
        
    Returns:
        Tuple of (retrieved documents, LLM response if available)
    """
    # Retrieve relevant documents
    documents = retriever.hybrid_retrieve(query, k=k)
    
    # Process with LLM if API key is available
    if "GOOGLE_API_KEY" in os.environ:
        gemini_processor = GeminiProcessor()
        response = gemini_processor.process_query(query, documents)
        return documents, response
    else:
        return documents, None

# Example usage:
retriever = process_document("knowledge-base/data/IT449_HR Roll back for policy Cancellation v2.1.docx")
results, answer = query_document(retriever, "What should HR system do?")
print("Retrieved documents:", results, "\nLLM answer:", answer)
